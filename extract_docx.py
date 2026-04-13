#!/usr/bin/env python3
"""
Extract plain text from DOCX files using python-docx
Usage: python extract_docx.py <path_to_docx>
"""
import sys
import os

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

def extract_text_from_docx(docx_path):
    """Extract all text from a DOCX file."""
    try:
        # Try to open the DOCX
        try:
            doc = Document(docx_path)
        except Exception as open_error:
            if "corrupt" in str(open_error).lower():
                print(f"ERROR: DOCX file is corrupted or invalid format", file=sys.stderr)
            else:
                print(f"ERROR: Could not open DOCX file: {str(open_error)}", file=sys.stderr)
            sys.exit(1)
        
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text.append(cell.text)
        
        # Check if we extracted any text
        if not text:
            print("ERROR: No extractable text found in DOCX. The file may be empty or contain only images.", file=sys.stderr)
            sys.exit(1)
        
        return "\n".join(text)
    except Exception as e:
        print(f"ERROR: Unexpected error: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("ERROR: No file path provided", file=sys.stderr)
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    if not os.path.exists(docx_path):
        print(f"ERROR: File not found: {docx_path}", file=sys.stderr)
        sys.exit(1)
    
    text = extract_text_from_docx(docx_path)
    print(text)
